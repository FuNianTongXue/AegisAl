from __future__ import annotations

import asyncio
import base64
import hashlib
import io
import re
import zipfile
from typing import Any, Literal
from xml.etree import ElementTree

from mcp.server.fastmcp import FastMCP
from pydantic import BaseModel, Field

from app.mcp.artifacts import MCPArtifactReference, stage_output_artifact


class WordReportOutput(BaseModel):
    schema_version: int = 1
    renderer: Literal["python-docx"] = "python-docx"
    artifacts: list[MCPArtifactReference] = Field(default_factory=list)
    media_type: Literal[
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_extension: Literal["docx"] = "docx"
    input_sha256: str
    output_sha256: str
    size: int
    page_preset: Literal["standard_business_brief"] = "standard_business_brief"
    header_template: Literal["memo_masthead"] = "memo_masthead"


report_word_mcp = FastMCP(
    "AegisAl Word MCP",
    instructions=(
        "Render verified AegisAl report JSON into a real DOCX using the standard_business_brief "
        "preset and memo_masthead header. Preserve headings, real lists, explicit table geometry, "
        "evidence line breaks, remediation, structured diagram summaries, and audit hashes."
    ),
)


@report_word_mcp.tool(
    name="render_word_report",
    description="Render verified AegisAl report JSON as a Microsoft Word DOCX artifact.",
    structured_output=True,
)
def render_word_report(
    report_document: dict[str, Any],
    mermaid: dict[str, Any] | None = None,
    *,
    output_dir: str,
) -> WordReportOutput:
    from app.reports import validate_report_document_json

    document = validate_report_document_json(report_document)
    payload = _build_docx(document, mermaid or {})
    if not payload.startswith(b"PK"):
        raise RuntimeError("Word MCP produced an invalid DOCX container")
    source = document.get("source") if isinstance(document.get("source"), dict) else {}
    input_sha256 = str((source.get("audit") or {}).get("payload_sha256") or "")
    digest = hashlib.sha256(payload).hexdigest()
    artifacts = [
        stage_output_artifact(
            output_dir,
            file_name="AegisAl-security-report.docx",
            payload=payload,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    ]
    return WordReportOutput(
        artifacts=artifacts,
        input_sha256=input_sha256,
        output_sha256=digest,
        size=len(payload),
    )


def invoke_report_word_mcp(arguments: dict[str, Any]) -> dict[str, Any]:
    result = asyncio.run(report_word_mcp.call_tool("render_word_report", arguments))
    if isinstance(result, tuple) and len(result) == 2 and isinstance(result[1], dict):
        return dict(result[1])
    raise RuntimeError("Word MCP did not return structured output")


async def word_mcp_spec() -> dict[str, Any]:
    tools = await report_word_mcp.list_tools()
    return {
        "id": "report-word",
        "name": report_word_mcp.name,
        "tools": [
            {
                "name": tool.name,
                "description": tool.description or "",
                "input_schema": tool.inputSchema,
                "output_schema": tool.outputSchema or {},
            }
            for tool in tools
        ],
    }


def _build_docx(document: dict[str, Any], mermaid: dict[str, Any]) -> bytes:
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is required for Word export") from exc

    report = document["report"]
    metadata = document.get("metadata") if isinstance(document.get("metadata"), dict) else {}
    template = document.get("template") if isinstance(document.get("template"), dict) else {}
    template_fonts = template.get("fonts") if isinstance(template.get("fonts"), dict) else {}
    template_tokens = template.get("style_tokens") if isinstance(template.get("style_tokens"), dict) else {}
    body_font = str(template_fonts.get("latin_body") or "SF Pro Text")
    heading_font = str(template_fonts.get("latin_heading") or "SF Pro Display")
    east_asia_font = str(template_fonts.get("body") or "PingFang SC")
    east_asia_heading_font = str(template_fonts.get("heading") or east_asia_font)
    code_font = str(template_fonts.get("code") or "SFMono-Regular")
    primary = _hex_token(template_tokens.get("primary"), "112C53")
    accent = _hex_token(template_tokens.get("accent"), "0BA3C4")
    text_color = _hex_token(template_tokens.get("text"), "15233A")
    from app.reports import _normalize_report_language, _report_china_time

    report_language = _normalize_report_language(
        metadata.get("language") or (report.get("metrics") or {}).get("language")
    )
    chinese_report = report_language in {"zh-Hans", "zh-Hant"}
    output = Document()
    section = output.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = output.styles
    _configure_style(styles["Normal"], body_font, east_asia_font, 10.5, RGBColor, qn, color=text_color, line_spacing=1.10, after=6)
    _configure_style(styles["Title"], heading_font, east_asia_heading_font, 23, RGBColor, qn, color=primary, bold=True, after=4)
    _configure_style(styles["Subtitle"], body_font, east_asia_font, 13, RGBColor, qn, color="52677F", after=12)
    heading_tokens = {
        "Heading 1": (16, primary, 16, 8),
        "Heading 2": (13, accent, 12, 6),
        "Heading 3": (11.5, primary, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        _configure_style(
            styles[name], heading_font, east_asia_heading_font, size, RGBColor, qn,
            color=color, bold=True, before=before, after=after, keep_with_next=True,
        )
    if "AegisAl Code" not in styles:
        code_style = styles.add_style("AegisAl Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["AegisAl Code"]
    _configure_style(code_style, code_font, east_asia_font, 8, RGBColor, qn, color="DCE6FF", after=0)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("AegisAl  /  安全扫描报告" if chinese_report else "AegisAl  /  SECURITY SCAN REPORT")
    _set_run_font(header_run, body_font, east_asia_font, qn, size=8, color=RGBColor(98, 112, 137), bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("AegisAl  |  ")
    _set_run_font(footer_run, body_font, east_asia_font, qn, size=8, color=RGBColor(114, 128, 150))
    _append_page_field(footer, OxmlElement, qn)

    title = output.add_paragraph(style="Title")
    title.add_run(str(report.get("project_name") or "AegisAl"))
    subtitle = output.add_paragraph(style="Subtitle")
    subtitle.add_run(str(report.get("title") or "安全扫描报告"))
    generated_at = _report_china_time(document.get("generated_at") or "-")
    source_hash = str((document.get("audit") or {}).get("source_payload_sha256") or "-")
    cover_rows = (
        (("生成时间", generated_at), ("来源 SHA-256", source_hash), ("状态", "扫描事实已核验"))
        if chinese_report
        else (("Generated", generated_at), ("Source SHA-256", source_hash), ("Status", "Verified scan facts"))
    )
    for label, value in cover_rows:
        paragraph = output.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        _set_run_font(label_run, body_font, east_asia_font, qn, size=9, bold=True)
        value_run = paragraph.add_run(value)
        _set_run_font(value_run, body_font, east_asia_font, qn, size=9)
    _paragraph_bottom_border(output.add_paragraph(), OxmlElement, qn, accent)

    metrics = report.get("metrics") if isinstance(report.get("metrics"), dict) else {}
    metric_table = output.add_table(rows=2, cols=4)
    metric_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    metric_table.autofit = False
    metric_values = [
        ("严重/高危", metrics.get("high_risk", 0)),
        ("中危", metrics.get("medium_risk", 0)),
        ("依赖漏洞", metrics.get("dependency_vulnerabilities", 0)),
        ("代码发现", metrics.get("code_findings", 0)),
    ]
    for column, (label, value) in enumerate(metric_values):
        metric_table.cell(0, column).text = str(value)
        metric_table.cell(1, column).text = label
    _set_table_geometry(metric_table, [2340, 2340, 2340, 2340], OxmlElement, qn)
    _style_table(metric_table, OxmlElement, qn, RGBColor, header_rows=0, centered=True)

    for index, item in enumerate(report.get("sections") or [], start=1):
        if not isinstance(item, dict):
            continue
        output.add_heading(f"{index}. {str(item.get('title') or '').strip()}", level=1)
        blocks = item.get("blocks") if isinstance(item.get("blocks"), list) else []
        if blocks:
            _append_report_blocks(output, blocks, document.get("visuals") or {}, styles)
        else:
            _append_markdown(
                output,
                str(item.get("content") or ""),
                styles,
                OxmlElement,
                qn,
                RGBColor,
                WD_ALIGN_PARAGRAPH,
                WD_CELL_VERTICAL_ALIGNMENT,
                WD_TABLE_ALIGNMENT,
            )

    core = output.core_properties
    core.title = str(report.get("title") or "神盾安全扫描报告")
    core.author = "AegisAl"
    core.subject = "Verified security scan report"
    core.keywords = "AegisAl, security, scan, report, MCP"
    _apply_document_run_fonts(
        output,
        qn,
        body_font=body_font,
        heading_font=heading_font,
        code_font=code_font,
        east_asia_font=east_asia_font,
    )
    buffer = io.BytesIO()
    output.save(buffer)
    return _add_ooxml_font_fallbacks(buffer.getvalue())


def _hex_token(value: Any, fallback: str) -> str:
    clean = str(value or "").strip().lstrip("#").upper()
    return clean if re.fullmatch(r"[0-9A-F]{6}", clean) else fallback


_WORDPROCESSINGML_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_OOXML_FONT_FALLBACKS = {
    "PingFang SC": "Arial Unicode MS",
    "SF Pro Display": "Arial",
    "SF Pro Text": "Arial",
    "SFMono-Regular": "Menlo",
}


def _add_ooxml_font_fallbacks(payload: bytes) -> bytes:
    """Keep the Apple font choices while giving non-Apple renderers safe aliases."""

    source = io.BytesIO(payload)
    target = io.BytesIO()
    with zipfile.ZipFile(source, mode="r") as archive, zipfile.ZipFile(
        target,
        mode="w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=6,
    ) as output:
        for entry in archive.infolist():
            content = archive.read(entry.filename)
            if entry.filename == "word/fontTable.xml":
                content = _patch_font_table(content)
            output.writestr(entry, content)
    return target.getvalue()


def _patch_font_table(content: bytes) -> bytes:
    namespace = _WORDPROCESSINGML_NAMESPACE
    ElementTree.register_namespace("w", namespace)
    root = ElementTree.fromstring(content)
    name_attribute = f"{{{namespace}}}name"
    value_attribute = f"{{{namespace}}}val"
    fonts = {
        str(node.get(name_attribute) or ""): node
        for node in root.findall(f"{{{namespace}}}font")
    }
    for preferred, fallback in _OOXML_FONT_FALLBACKS.items():
        font = fonts.get(preferred)
        if font is None:
            font = ElementTree.SubElement(root, f"{{{namespace}}}font", {name_attribute: preferred})
        alternative = font.find(f"{{{namespace}}}altName")
        if alternative is None:
            alternative = ElementTree.SubElement(font, f"{{{namespace}}}altName")
        alternative.set(value_attribute, fallback)
        family = font.find(f"{{{namespace}}}family")
        if family is None:
            family = ElementTree.SubElement(font, f"{{{namespace}}}family")
        family.set(value_attribute, "modern" if preferred == "SFMono-Regular" else "swiss")
        pitch = font.find(f"{{{namespace}}}pitch")
        if pitch is None:
            pitch = ElementTree.SubElement(font, f"{{{namespace}}}pitch")
        pitch.set(value_attribute, "fixed" if preferred == "SFMono-Regular" else "variable")
    return ElementTree.tostring(root, encoding="UTF-8", xml_declaration=True)


def _append_markdown(
    document: Any,
    markdown: str,
    styles: Any,
    OxmlElement: Any,
    qn: Any,
    RGBColor: Any,
    WD_ALIGN_PARAGRAPH: Any,
    WD_CELL_VERTICAL_ALIGNMENT: Any,
    WD_TABLE_ALIGNMENT: Any,
) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        image_match = re.fullmatch(
            r"!\[[^\]]*\]\(data:image/(png|jpeg);base64,([A-Za-z0-9+/=]+)\)",
            stripped,
        )
        if image_match:
            from docx.shared import Inches

            try:
                payload = base64.b64decode(image_match.group(2), validate=True)
            except (TypeError, ValueError) as exc:
                raise ValueError("Word report contains an invalid Mermaid image") from exc
            valid_signature = (
                payload.startswith(b"\x89PNG\r\n\x1a\n")
                if image_match.group(1) == "png"
                else payload.startswith(b"\xff\xd8\xff")
            )
            if len(payload) > 8 * 1024 * 1024 or not valid_signature:
                raise ValueError("Word report Mermaid image failed validation")
            document.add_picture(io.BytesIO(payload), width=Inches(6.1))
            index += 1
            continue
        if stripped.startswith("<!-- secflow-mermaid-source:"):
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            _add_code_block(document, code_lines, language, styles, OxmlElement, qn, RGBColor)
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(
                document, table_lines, OxmlElement, qn, RGBColor,
                WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT,
            )
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            document.add_heading(_plain_markdown(heading.group(2)), level=min(3, len(heading.group(1)) - 1))
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(_plain_markdown(stripped[2:]))
        elif re.match(r"^\d+[.)]\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(_plain_markdown(re.sub(r"^\d+[.)]\s+", "", stripped)))
        elif stripped.startswith("> "):
            from docx.shared import Inches

            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(_plain_markdown(stripped[2:]))
            run.italic = True
            _shade_paragraph(paragraph, OxmlElement, qn, "EDF7FA")
        elif stripped != "---" and not stripped.startswith("<!--"):
            document.add_paragraph(_plain_markdown(stripped))
        index += 1


def _append_report_blocks(
    document: Any,
    blocks: list[dict[str, Any]],
    visuals: dict[str, Any],
    styles: Any,
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, RGBColor

    diagrams = {
        str(item.get("image_sha256") or ""): item
        for item in visuals.get("diagrams") or []
        if isinstance(item, dict)
    }
    for block in blocks:
        if not isinstance(block, dict):
            continue
        kind = str(block.get("type") or "")
        if kind == "heading":
            document.add_heading(str(block.get("text") or ""), level=min(3, max(1, int(block.get("level") or 2))))
        elif kind == "paragraph":
            document.add_paragraph(str(block.get("text") or ""))
        elif kind == "quote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(str(block.get("text") or ""))
            run.italic = True
            _shade_paragraph(paragraph, OxmlElement, qn, "EDF7FA")
        elif kind in {"bullet_list", "numbered_list"}:
            style = "List Bullet" if kind == "bullet_list" else "List Number"
            for item in block.get("items") or []:
                document.add_paragraph(str(item), style=style)
        elif kind == "code":
            _add_code_block(
                document,
                [str(line) for line in block.get("lines") or []],
                str(block.get("language") or ""),
                styles,
                OxmlElement,
                qn,
                RGBColor,
            )
        elif kind == "table":
            _add_json_table(
                document,
                [str(item) for item in block.get("columns") or []],
                [[str(cell) for cell in row] for row in block.get("rows") or [] if isinstance(row, list)],
                OxmlElement,
                qn,
                RGBColor,
                WD_CELL_VERTICAL_ALIGNMENT,
                WD_TABLE_ALIGNMENT,
            )
        elif kind == "diagram":
            diagram = diagrams.get(str(block.get("sha256") or ""))
            if not diagram:
                raise ValueError("Word report JSON references a missing diagram")
            payload = base64.b64decode(str(diagram.get("image_base64") or ""), validate=True)
            media_type = str(diagram.get("image_media_type") or "")
            valid = (media_type == "image/jpeg" and payload.startswith(b"\xff\xd8\xff")) or (
                media_type == "image/png" and payload.startswith(b"\x89PNG\r\n\x1a\n")
            )
            if not valid or hashlib.sha256(payload).hexdigest() != str(diagram.get("image_sha256") or ""):
                raise ValueError("Word report JSON diagram failed validation")
            try:
                from PIL import Image as PillowImage

                with PillowImage.open(io.BytesIO(payload)) as bitmap:
                    pixel_width, pixel_height = bitmap.size
            except Exception as exc:  # noqa: BLE001
                raise ValueError("Word report JSON diagram dimensions are invalid") from exc
            scale = min(6.1 / max(1, pixel_width), 8.0 / max(1, pixel_height))
            document.add_picture(
                io.BytesIO(payload),
                width=Inches(max(0.1, pixel_width * scale)),
                height=Inches(max(0.1, pixel_height * scale)),
            )


def _add_json_table(
    document: Any,
    columns: list[str],
    rows: list[list[str]],
    OxmlElement: Any,
    qn: Any,
    RGBColor: Any,
    WD_CELL_VERTICAL_ALIGNMENT: Any,
    WD_TABLE_ALIGNMENT: Any,
) -> None:
    column_count = max(1, len(columns), *(len(row) for row in rows))
    table = document.add_table(rows=len(rows) + 1, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for column_index in range(column_count):
        table.cell(0, column_index).text = columns[column_index] if column_index < len(columns) else ""
    for row_index, row in enumerate(rows, start=1):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = row[column_index] if column_index < len(row) else ""
    if column_count == 2:
        widths = [2100, 7260]
    elif column_count == 8:
        widths = [600, 800, 1300, 1000, 1400, 1400, 1860, 1000]
    elif column_count == 9:
        widths = [850, 650, 650, 900, 850, 1050, 1150, 850, 2410]
    else:
        base = 9360 // column_count
        widths = [base] * column_count
        widths[-1] += 9360 - sum(widths)
    _set_table_geometry(table, widths, OxmlElement, qn)
    _style_table(table, OxmlElement, qn, RGBColor, header_rows=1, centered=False)


def _add_markdown_table(
    document: Any,
    lines: list[str],
    OxmlElement: Any,
    qn: Any,
    RGBColor: Any,
    WD_CELL_VERTICAL_ALIGNMENT: Any,
    WD_TABLE_ALIGNMENT: Any,
) -> None:
    rows = [_split_table_row(line) for line in lines]
    rows = [row for row in rows if row and not all(set(cell) <= {"-", ":", " "} for cell in row)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = _plain_markdown(row[column_index] if column_index < len(row) else "")
    if column_count == 2:
        widths = [2100, 7260]
    elif column_count == 8:
        widths = [600, 800, 1300, 1000, 1400, 1400, 1860, 1000]
    elif column_count == 9:
        widths = [850, 650, 650, 900, 850, 1050, 1150, 850, 2410]
    else:
        base = 9360 // column_count
        widths = [base] * column_count
        widths[-1] += 9360 - sum(widths)
    _set_table_geometry(table, widths, OxmlElement, qn)
    _style_table(table, OxmlElement, qn, RGBColor, header_rows=1, centered=False)


def _add_code_block(
    document: Any,
    lines: list[str],
    language: str,
    styles: Any,
    OxmlElement: Any,
    qn: Any,
    RGBColor: Any,
) -> None:
    if language.lower() == "mermaid":
        from app.reports import _mermaid_structured_rows

        headers, rows = _mermaid_structured_rows("\n".join(lines))
        if not rows:
            return
        table = document.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = 0
        table.autofit = False
        for column, value in enumerate(headers):
            table.cell(0, column).text = value
        for row_index, row in enumerate(rows, start=1):
            for column, value in enumerate(row):
                table.cell(row_index, column).text = value
        widths = [3100, 1800, 4460] if len(headers) == 3 else [6960, 2400]
        _set_table_geometry(table, widths, OxmlElement, qn)
        _style_table(table, OxmlElement, qn, RGBColor, header_rows=1, centered=False)
        return
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_geometry(table, [9360], OxmlElement, qn)
    cell = table.cell(0, 0)
    _shade_cell(cell, OxmlElement, qn, "111936")
    paragraph = cell.paragraphs[0]
    paragraph.style = styles["AegisAl Code"]
    for offset, line in enumerate(lines):
        if offset:
            paragraph.add_run().add_break()
        paragraph.add_run(line or " ")


def _configure_style(
    style: Any,
    latin: str,
    east_asia: str,
    size: float,
    RGBColor: Any,
    qn: Any,
    *,
    color: str = "26364D",
    bold: bool = False,
    before: float = 0,
    after: float = 0,
    line_spacing: float | None = None,
    keep_with_next: bool = False,
) -> None:
    from docx.shared import Pt

    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    _set_rfonts(style._element.rPr.rFonts, latin, east_asia, qn)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.keep_with_next = keep_with_next


def _set_run_font(
    run: Any,
    latin: str,
    east_asia: str,
    qn: Any,
    *,
    size: float,
    color: Any | None = None,
    bold: bool = False,
) -> None:
    from docx.shared import Pt

    run.font.name = latin
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    _set_rfonts(run._element.get_or_add_rPr().rFonts, latin, east_asia, qn)


def _set_rfonts(rfonts: Any, latin: str, east_asia: str, qn: Any) -> None:
    for theme_name in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(qn(f"w:{theme_name}"), None)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def _apply_document_run_fonts(
    document: Any,
    qn: Any,
    *,
    body_font: str = "SF Pro Text",
    heading_font: str = "SF Pro Display",
    code_font: str = "SFMono-Regular",
    east_asia_font: str = "PingFang SC",
) -> None:
    from docx.oxml import OxmlElement

    def paragraphs(container: Any):
        yield from container.paragraphs
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from paragraphs(cell)

    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.footer])
    for container in containers:
        for paragraph in paragraphs(container):
            style_name = str(getattr(paragraph.style, "name", "") or "")
            if style_name == "AegisAl Code":
                latin = code_font
            elif style_name in {"Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"}:
                latin = heading_font
            else:
                latin = body_font
            for run in paragraph.runs:
                rpr = run._element.get_or_add_rPr()
                contains_cjk = bool(re.search(r"[\u2e80-\u9fff\uf900-\ufaff]", str(run.text or "")))
                run_latin = "Arial Unicode MS" if contains_cjk else latin
                rfonts = rpr.get_or_add_rFonts()
                _set_rfonts(rfonts, run_latin, east_asia_font, qn)
                if contains_cjk:
                    rfonts.set(qn("w:hint"), "eastAsia")
                language = rpr.find(qn("w:lang"))
                if language is None:
                    language = OxmlElement("w:lang")
                    rpr.append(language)
                language.set(qn("w:val"), "en-US")
                language.set(qn("w:eastAsia"), "zh-CN")


def _set_table_geometry(table: Any, widths: list[int], OxmlElement: Any, qn: Any) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell_width = widths[min(index, len(widths) - 1)]
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(cell_width))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell, OxmlElement, qn)


def _set_cell_margins(cell: Any, OxmlElement: Any, qn: Any) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _style_table(
    table: Any,
    OxmlElement: Any,
    qn: Any,
    RGBColor: Any,
    *,
    header_rows: int,
    centered: bool,
) -> None:
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        if row_index < header_rows and row_properties.find(qn("w:tblHeader")) is None:
            row_properties.append(OxmlElement("w:tblHeader"))
        for cell in row.cells:
            if row_index < header_rows:
                _shade_cell(cell, OxmlElement, qn, "F2F4F7")
            for paragraph in cell.paragraphs:
                if centered:
                    paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.size = __import__("docx").shared.Pt(9)
                    if row_index < header_rows:
                        run.bold = True
                    run.font.color.rgb = RGBColor(38, 54, 77)


def _shade_cell(cell: Any, OxmlElement: Any, qn: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _shade_paragraph(paragraph: Any, OxmlElement: Any, qn: Any, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _paragraph_bottom_border(paragraph: Any, OxmlElement: Any, qn: Any, color: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    properties.append(borders)


def _append_page_field(paragraph: Any, OxmlElement: Any, qn: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _split_table_row(value: str) -> list[str]:
    return [cell.strip().replace("\\|", "|").replace("<br>", "\n") for cell in value.strip().strip("|").split("|")]


def _plain_markdown(value: str) -> str:
    clean = re.sub(r"!\[([^]]*)\]\([^)]+\)", r"\1", str(value or ""))
    clean = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", clean)
    clean = re.sub(r"`([^`]+)`", r"\1", clean)
    clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", clean)
    return clean.replace("<br>", "\n").strip()


def main() -> None:
    report_word_mcp.run(transport="stdio")


if __name__ == "__main__":
    main()
